"""
scripts/proposal_generator.py
-----------------------------
Minimal DOCX proposal builder used by BidForge / RFP Auto-Respond when
WeasyPrint is unavailable (typical on Windows without GTK/Pango).

Public API expected by the rest of the codebase:

  generate(cfg, output_docx_path) -> str
  add_section(doc, cfg, section, index=1) -> None
  convert_to_pdf(docx_path, output_dir) -> Optional[str]
"""

from __future__ import annotations

import logging
import os
import re
import shutil
import time
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    _DOCX_OK = True
except ImportError:
    Document = None  # type: ignore
    WD_ALIGN_PARAGRAPH = None  # type: ignore
    OxmlElement = None  # type: ignore
    qn = None  # type: ignore
    Inches = None  # type: ignore
    Pt = None  # type: ignore
    RGBColor = None  # type: ignore
    _DOCX_OK = False


# ---------------------------------------------------------------------------
# Inline markdown (**bold**, *italic*, `code`) → runs
# ---------------------------------------------------------------------------

_INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|[^*`]+)"
)


def _add_runs_with_inline(paragraph, text: str) -> None:
    if not text:
        return
    for token in _INLINE_RE.findall(text):
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(token)


def _set_cell_shading(cell, hex_color: str) -> None:
    """Set table header cell background color."""
    color = hex_color.lstrip("#")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _accent_rgb(brand: dict[str, Any]) -> RGBColor:
    raw = str(brand.get("accent_color") or "2B6CB0").strip().lstrip("#")
    try:
        return RGBColor(int(raw[0:2], 16), int(raw[2:4], 16), int(raw[4:6], 16))
    except Exception:
        return RGBColor(0x2B, 0x6C, 0xB0)


# ---------------------------------------------------------------------------
# Section writer
# ---------------------------------------------------------------------------

def add_section(doc, cfg: dict[str, Any], section: dict[str, Any], index: int = 1) -> None:
    """Append one outline section (title + blocks) to an open Document."""
    if not _DOCX_OK:
        raise RuntimeError("python-docx is required for proposal generation")

    brand = cfg.get("brand") or {}
    title = (section.get("title") or f"Section {index}").strip()
    blocks: list[dict[str, Any]] = section.get("blocks") or []

    if section.get("page_break_before"):
        doc.add_page_break()

    # Section heading
    try:
        heading = doc.add_heading(title, level=1)
    except Exception:
        heading = doc.add_paragraph(title)
        for run in heading.runs:
            run.bold = True
            run.font.size = Pt(16)
    try:
        for run in heading.runs:
            run.font.color.rgb = _accent_rgb(brand)
    except Exception:
        pass

    for block in blocks:
        btype = (block.get("type") or "paragraph").lower()

        if btype == "paragraph":
            p = doc.add_paragraph()
            _add_runs_with_inline(p, str(block.get("text") or ""))

        elif btype == "subheading":
            level = min(max(int(block.get("level") or 2), 2), 4)
            text = str(block.get("text") or "")
            try:
                doc.add_heading(text, level=level)
            except Exception:
                p = doc.add_paragraph(text)
                for run in p.runs:
                    run.bold = True

        elif btype == "bullets":
            for item in block.get("items") or []:
                try:
                    p = doc.add_paragraph(style="List Bullet")
                except Exception:
                    p = doc.add_paragraph(style=None)
                    p.add_run("• ")
                _add_runs_with_inline(p, str(item))

        elif btype == "numbered":
            for item in block.get("items") or []:
                try:
                    p = doc.add_paragraph(style="List Number")
                except Exception:
                    p = doc.add_paragraph()
                _add_runs_with_inline(p, str(item))

        elif btype == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(str(block.get("text") or ""))
            run.italic = True

        elif btype == "divider":
            p = doc.add_paragraph("─" * 40)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

        elif btype == "table":
            headers = block.get("headers") or []
            rows = block.get("rows") or []
            cols = max(len(headers), max((len(r) for r in rows), default=0), 1)
            table = doc.add_table(rows=1 + len(rows), cols=cols)
            try:
                table.style = "Table Grid"
            except Exception:
                pass
            # Header row
            for i, h in enumerate(headers):
                if i >= cols:
                    break
                cell = table.rows[0].cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                clean_h = re.sub(r"^\*{1,2}(.*?)\*{1,2}$", r"\1", str(h).strip())
                _add_runs_with_inline(p, clean_h)
                for run in p.runs:
                    run.bold = True
                try:
                    _set_cell_shading(cell, str(brand.get("accent_color") or "2B6CB0"))
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                except Exception:
                    pass
            # Data rows
            for r_idx, row in enumerate(rows):
                for c_idx in range(cols):
                    val = row[c_idx] if c_idx < len(row) else ""
                    cell = table.rows[r_idx + 1].cells[c_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    _add_runs_with_inline(p, str(val).strip())
            doc.add_paragraph()

        else:
            # Unknown block type — dump text if present
            text = block.get("text")
            if text:
                p = doc.add_paragraph()
                _add_runs_with_inline(p, str(text))


def generate(cfg: dict[str, Any], output_docx_path: str) -> str:
    """Build a full .docx from cfg {brand, proposal, sections}."""
    if not _DOCX_OK:
        raise RuntimeError("python-docx is required. pip install python-docx")

    brand = cfg.get("brand") or {}
    proposal = cfg.get("proposal") or {}
    sections: list[dict[str, Any]] = cfg.get("sections") or []

    doc = Document()

    # Cover-ish header
    title = proposal.get("title") or "Proposal"
    subtitle = proposal.get("subtitle") or ""
    prepared_for = proposal.get("prepared_for") or ""
    prepared_by = proposal.get("prepared_by") or "OrbitAvanya Tech LLP"
    proposal_date = proposal.get("proposal_date") or ""

    h = doc.add_heading(title, level=0)
    try:
        for run in h.runs:
            run.font.color.rgb = _accent_rgb(brand)
    except Exception:
        pass

    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.italic = True

    meta_lines = []
    if prepared_for:
        meta_lines.append(f"Prepared for: {prepared_for}")
    if prepared_by:
        meta_lines.append(f"Prepared by: {prepared_by}")
    if proposal_date:
        meta_lines.append(f"Date: {proposal_date}")
    if proposal.get("engagement_ref"):
        meta_lines.append(f"Ref: {proposal['engagement_ref']}")
    if meta_lines:
        p = doc.add_paragraph("\n".join(meta_lines))

    conf = proposal.get("confidentiality_text")
    if conf:
        p = doc.add_paragraph()
        run = p.add_run(str(conf))
        run.italic = True
        run.font.size = Pt(9)

    doc.add_page_break()

    for i, section in enumerate(sections, start=1):
        # First body section already follows a page break from the cover.
        sec = dict(section)
        if i == 1:
            sec["page_break_before"] = False
        add_section(doc, cfg, sec, index=i)

    out = Path(output_docx_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    # Save via temp file so a locked destination (browser/Word open) does not
    # crash generation on Windows PermissionError.
    tmp = out.with_suffix(f".tmp.{os.getpid()}.docx")
    try:
        doc.save(str(tmp))
        try:
            if out.exists():
                out.unlink()
            tmp.replace(out)
        except OSError:
            alt = out.with_name(f"{out.stem}_{int(time.time())}{out.suffix}")
            tmp.replace(alt)
            logger.warning(f"[proposal_generator] Dest locked; wrote {alt}")
            return str(alt)
    except Exception:
        # Fallback: direct save
        doc.save(str(out))
    logger.info(f"[proposal_generator] Wrote DOCX: {out}")
    return str(out)


def convert_to_pdf(docx_path: str, output_dir: str | None = None) -> str | None:
    """
    Best-effort DOCX → PDF.

    Tries, in order:
      1. docx2pdf (Windows + Word)
      2. LibreOffice headless
    Returns PDF path on success, or the original DOCX path if conversion
    is unavailable (so the pipeline still has a downloadable artifact).
    """
    src = Path(docx_path)
    out_dir = Path(output_dir) if output_dir else src.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = out_dir / (src.stem + ".pdf")

    # 1) docx2pdf
    try:
        from docx2pdf import convert as docx2pdf_convert
        docx2pdf_convert(str(src), str(pdf_path))
        if pdf_path.exists():
            logger.info(f"[proposal_generator] PDF via docx2pdf: {pdf_path}")
            return str(pdf_path)
    except Exception as exc:
        logger.debug(f"[proposal_generator] docx2pdf unavailable: {exc}")

    # 2) LibreOffice / soffice
    for binary in ("soffice", "libreoffice"):
        try:
            import subprocess
            result = subprocess.run(
                [binary, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(src)],
                capture_output=True,
                text=True,
                timeout=120,
            )
            if pdf_path.exists():
                logger.info(f"[proposal_generator] PDF via {binary}: {pdf_path}")
                return str(pdf_path)
            logger.debug(f"[proposal_generator] {binary} exit={result.returncode} stderr={result.stderr[:200]}")
        except Exception as exc:
            logger.debug(f"[proposal_generator] {binary} unavailable: {exc}")

    # No PDF converter — ship the DOCX (UI accepts .docx downloads).
    logger.warning(
        "[proposal_generator] PDF conversion unavailable; returning DOCX path. "
        "Install Microsoft Word + docx2pdf, or LibreOffice, for PDF output."
    )
    # Ensure a copy sits in output_dir if needed
    dest = out_dir / src.name
    if src.resolve() != dest.resolve():
        try:
            shutil.copy2(src, dest)
            return str(dest)
        except Exception:
            pass
    return str(src)
