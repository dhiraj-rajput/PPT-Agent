"""
documents/bidforge/first_page_preserver.py
--------------------------------------------
Preserves the *actual* first page of an uploaded .docx template (cover page,
registration/company details, etc.) exactly as the user wrote it, and appends
the AI-generated proposal body after it -- instead of the old approach of
discarding the template entirely and generating a brand-new cover page from
scratch.

WHAT "FIRST PAGE" MEANS HERE
----------------------------
python-docx (and OOXML in general) has no native concept of a rendered page --
page breaks are a layout-engine concern. We approximate "first page" as
everything before the first *explicit* page break (a `<w:br w:type="page"/>`
run) in the document body. If the template has no explicit page break at all,
the entire template is treated as the first page and new content is appended
after it on a fresh page.

WHAT GETS CHANGED VS. LEFT ALONE
---------------------------------
- Only text that matches a recognized date pattern (e.g. "March 3, 2025",
  "03/03/2025", "2025-03-03") AND falls within a single run's text AND
  does not equal today's date gets rewritten -- to today's date, in the same
  format it was already written in.
- Everything else on the preserved page (contact info, names, addresses,
  registration numbers, logos, tables) is left byte-for-byte untouched.
- We only ever operate on a copy of the uploaded template, never the
  original upload path.
"""

from __future__ import annotations

import re
import shutil
import tempfile
from datetime import date
from pathlib import Path
from typing import Any, Dict, Iterator, List, Optional

from utils.helpers import setup_logger

logger = setup_logger(__name__)

_PAGE_BREAK_QN = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"
_PAGE_BREAK_TYPE_QN = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"


# ---------------------------------------------------------------------------
# First-page detection
# ---------------------------------------------------------------------------

def _paragraph_has_page_break(paragraph) -> bool:
    """True if this paragraph contains an explicit manual page break run."""
    for br in paragraph._p.findall(f".//{_PAGE_BREAK_QN}"):
        if br.get(_PAGE_BREAK_TYPE_QN) == "page":
            return True
    return False


_PAGE_BREAK_BEFORE_QN = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pageBreakBefore"

# Canonical section names commonly used as the FIRST body heading in company
# templates. If the template's own first heading-styled paragraph matches
# one of these (case-insensitively, ignoring leading numbering like "1. "),
# that is treated as the start of body content -- everything from there
# onward is NOT preserved, because the AI-generated body will write its own
# version of this exact section. This is what previously caused duplicate
# "Executive Summary" sections: a template with no *explicit* page break but
# with its own Executive Summary heading a page or two in was being kept in
# full, and the AI then wrote a second one after it.
_BODY_START_HEADING_PATTERNS = [
    "executive summary", "introduction", "scope of work", "about us",
    "company profile", "company overview", "our understanding",
    "understanding of requirements", "proposed solution", "technical approach",
    "cover letter", "table of contents",
]

_HEADING_STYLE_PREFIXES = ("heading", "title")


def _paragraph_style_name(paragraph_el) -> str:
    """Best-effort style name lookup without needing a bound Paragraph object
    (we're walking raw body children, not doc.paragraphs, so we read the
    pStyle w:val directly)."""
    pPr = paragraph_el.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
    if pPr is None:
        return ""
    pStyle = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle")
    if pStyle is None:
        return ""
    return pStyle.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") or ""


def _paragraph_text(paragraph_el) -> str:
    return "".join(node.text or "" for node in paragraph_el.iter() if node.tag.endswith("}t")).strip()


def _normalize_heading(text: str) -> str:
    """Strip leading numbering ('1.', '1)', 'Section 1:') and punctuation so
    '1. Executive Summary' and 'EXECUTIVE SUMMARY' compare equal."""
    t = re.sub(r"^\s*(section\s+)?\d+[\.\)\:]?\s*", "", text.strip(), flags=re.IGNORECASE)
    t = re.sub(r"[^a-z0-9 ]", "", t.lower()).strip()
    t = re.sub(r"\s+", " ", t)
    return t


def _looks_like_body_start_heading(text: str) -> bool:
    norm = _normalize_heading(text)
    if not norm:
        return False
    return any(norm == pat or norm.startswith(pat) for pat in _BODY_START_HEADING_PATTERNS)


def find_first_page_split_index(doc) -> Optional[int]:
    """Return the index (into doc.element.body's <w:p>/<w:tbl> children) of
    the paragraph where the template's true cover/front-matter ends, or None
    if the whole template should be treated as front matter (no signal
    found at all).

    Three signals are checked, in document order, and the EARLIEST one wins
    (whichever indicates body content starts first):
      1. An explicit manual page break run (<w:br w:type="page"/>).
      2. A paragraph with the pageBreakBefore paragraph property set.
      3. The first heading-styled paragraph (style name starting with
         "Heading"/"Title") whose text matches a known first-body-section
         name (Executive Summary, Introduction, Scope of Work, etc.) -- this
         catches templates that flow straight from cover page into body
         content with no explicit break at all, which is the common case
         that previously caused the whole template (including its own
         Executive Summary) to be preserved wholesale.
    """
    body = doc.element.body
    candidates: List[int] = []
    for idx, child in enumerate(body):
        tag = child.tag.rsplit("}", 1)[-1]
        if tag != "p":
            continue
        for br in child.findall(f".//{_PAGE_BREAK_QN}"):
            if br.get(_PAGE_BREAK_TYPE_QN) == "page":
                candidates.append(idx)
                break
        pPr = child.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
        if pPr is not None and pPr.find(_PAGE_BREAK_BEFORE_QN) is not None and idx > 0:
            candidates.append(idx - 1)  # preserve up to (not including) this paragraph
        if idx > 0:
            style = _paragraph_style_name(child).lower()
            if any(style.startswith(p) for p in _HEADING_STYLE_PREFIXES):
                text = _paragraph_text(child)
                if _looks_like_body_start_heading(text):
                    candidates.append(idx - 1)
    if not candidates:
        return None
    return min(candidates)


def detect_preserved_headings(doc, split_idx: Optional[int]) -> List[str]:
    """Returns normalized heading text for every heading-styled paragraph
    that WILL be kept in the preserved region (index <= split_idx, or the
    whole document if split_idx is None). Used by document_generator.py to
    tell the section-writer prompt which section titles already exist in the
    template so it doesn't author a second copy of them."""
    body = doc.element.body
    headings: List[str] = []
    for idx, child in enumerate(body):
        if split_idx is not None and idx > split_idx:
            break
        tag = child.tag.rsplit("}", 1)[-1]
        if tag != "p":
            continue
        style = _paragraph_style_name(child).lower()
        if any(style.startswith(p) for p in _HEADING_STYLE_PREFIXES):
            text = _paragraph_text(child)
            if text:
                headings.append(text)
    return headings


def iter_first_page_paragraph_texts(doc) -> Iterator[str]:
    """Yield paragraph text for everything on the template's first page
    (or the whole document if no page break is found). Used for both
    AI-context summarization and as a light-weight preview."""
    split_idx = find_first_page_split_index(doc)
    body = doc.element.body
    for idx, child in enumerate(body):
        if split_idx is not None and idx > split_idx:
            break
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            text = "".join(node.text or "" for node in child.iter() if node.tag.endswith("}t"))
            if text:
                yield text


# ---------------------------------------------------------------------------
# Stale-date detection & in-place replacement (single-run matches only)
# ---------------------------------------------------------------------------

_MONTHS = [
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
]
_MONTH_RE = "|".join(_MONTHS)
_MONTH_INDEX = {m.lower(): i + 1 for i, m in enumerate(_MONTHS)}

_DATE_PATTERNS = [
    # "March 3, 2025" / "March 3rd 2025"
    re.compile(rf"\b({_MONTH_RE})\s+(\d{{1,2}})(st|nd|rd|th)?,?\s+(\d{{4}})\b", re.IGNORECASE),
    # "3 March 2025"
    re.compile(rf"\b(\d{{1,2}})(st|nd|rd|th)?\s+({_MONTH_RE})\s+(\d{{4}})\b", re.IGNORECASE),
    # "03/03/2025" or "03-03-2025"
    re.compile(r"\b(\d{1,2})([/\-])(\d{1,2})\2(\d{4})\b"),
    # ISO "2025-03-03"
    re.compile(r"\b(\d{4})-(\d{2})-(\d{2})\b"),
]


def _safe_date(year: int, month: int, day: int) -> Optional[date]:
    try:
        return date(year, month, day)
    except ValueError:
        return None


def _parse_matched_date(match: "re.Match", pattern_index: int) -> Optional[date]:
    groups = match.groups()
    try:
        if pattern_index == 0:  # Month D, YYYY
            month = _MONTH_INDEX[groups[0].lower()]
            return _safe_date(int(groups[3]), month, int(groups[1]))
        if pattern_index == 1:  # D Month YYYY
            month = _MONTH_INDEX[groups[2].lower()]
            return _safe_date(int(groups[3]), month, int(groups[0]))
        if pattern_index == 2:  # numeric MM/DD/YYYY (US-style, matches template convention)
            return _safe_date(int(groups[3]), int(groups[0]), int(groups[2]))
        if pattern_index == 3:  # ISO YYYY-MM-DD
            return _safe_date(int(groups[0]), int(groups[1]), int(groups[2]))
    except (KeyError, ValueError):
        return None
    return None


def _format_like(pattern_index: int, match: "re.Match", today: date) -> str:
    groups = match.groups()
    if pattern_index == 0:  # "Month D, YYYY" (keep month name + optional ordinal + comma style)
        month_name = _MONTHS[today.month - 1]
        suffix = groups[2] or ""
        comma = "," if "," in match.group(0) else ""
        return f"{month_name} {today.day}{suffix}{comma} {today.year}"
    if pattern_index == 1:  # "D Month YYYY"
        month_name = _MONTHS[today.month - 1]
        suffix = groups[1] or ""
        return f"{today.day}{suffix} {month_name} {today.year}"
    if pattern_index == 2:  # numeric, keep original separator
        sep = groups[1]
        return f"{today.month:02d}{sep}{today.day:02d}{sep}{today.year:04d}"
    if pattern_index == 3:  # ISO
        return f"{today.year:04d}-{today.month:02d}-{today.day:02d}"
    return match.group(0)


def _update_stale_dates_in_text(text: str, today: date) -> str:
    """Replace only dates that (a) match a recognized format and (b) don't
    already equal today, with today's date in the same format. Anything that
    doesn't parse as a confident date (phone numbers, IDs, etc.) is left
    completely alone."""
    if not text or not any(ch.isdigit() for ch in text):
        return text

    for pattern_index, pattern in enumerate(_DATE_PATTERNS):
        def _replacer(match: "re.Match") -> str:
            parsed = _parse_matched_date(match, pattern_index)
            if parsed is None or parsed == today:
                return match.group(0)
            return _format_like(pattern_index, match, today)

        text = pattern.sub(_replacer, text)
    return text


def _update_dates_in_preserved_region(doc, split_idx: Optional[int], today: Optional[date] = None) -> int:
    """Walk every paragraph in the preserved first-page region and patch
    stale dates in place, run by run. Only touches a run when the *entire*
    matched date lives inside that single run's text, so formatting/runs that
    also carry unrelated contact info are never split or corrupted.
    Returns the number of runs changed."""
    today = today or date.today()
    body = doc.element.body
    changed = 0
    for idx, child in enumerate(body):
        if split_idx is not None and idx > split_idx:
            break
        tag = child.tag.rsplit("}", 1)[-1]
        if tag != "p":
            continue
        for run_el in child.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"):
            for t_el in run_el.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                original = t_el.text or ""
                updated = _update_stale_dates_in_text(original, today)
                if updated != original:
                    t_el.text = updated
                    changed += 1
    return changed


# ---------------------------------------------------------------------------
# Style bootstrap -- appended content uses "Heading 1" / "List Bullet" /
# "List Number" / "Table Grid" styles; make sure the (possibly minimal)
# uploaded template actually has them before we lean on proposal_generator's
# block handlers, which assume those built-ins exist.
# ---------------------------------------------------------------------------

_REQUIRED_STYLES = ["Heading 1", "List Bullet", "List Number", "Table Grid"]


def _ensure_required_styles(doc) -> None:
    from docx import Document as _Document

    missing = [name for name in _REQUIRED_STYLES if name not in {s.name for s in doc.styles}]
    if not missing:
        return
    try:
        blank = _Document()
        existing_ids = {el.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId") for el in doc.styles.element}
        for style in blank.styles:
            if style.name not in missing:
                continue
            style_id = style.style_id
            if style_id in existing_ids:
                continue
            doc.styles.element.append(style.element)
    except Exception as e:
        logger.warning(f"[FirstPagePreserver] Could not backfill missing styles {missing}: {e}")


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def get_preserved_headings(template_path: str) -> List[str]:
    """Preview-only helper: opens the template read-only and returns the
    heading titles that would be preserved verbatim, WITHOUT writing
    anything. document_generator.py calls this before running any section
    -writer LLM calls, so it can tell each call "this heading already exists
    in the template, don't write it again" up front instead of generating
    duplicate content and trying to strip it out afterward."""
    from docx import Document
    try:
        doc = Document(template_path)
    except Exception as exc:
        logger.warning(f"[FirstPagePreserver] Could not open template to preview headings: {exc}")
        return []
    split_idx = find_first_page_split_index(doc)
    return detect_preserved_headings(doc, split_idx)


def dedupe_sections_against_preserved(
    sections: List[Dict[str, Any]], preserved_headings: List[str]
) -> List[Dict[str, Any]]:
    """Drops any outline/generated section whose title normalizes to the same
    thing as a heading already present in the preserved template region.
    Belt-and-suspenders alongside telling the LLM up front -- if the model
    ignores the instruction and writes the section anyway, this still keeps
    it out of the final document instead of shipping a duplicate."""
    if not preserved_headings:
        return sections
    preserved_norm = {_normalize_heading(h) for h in preserved_headings}
    kept = []
    for s in sections:
        title = str(s.get("title") or s.get("key") or "")
        if _normalize_heading(title) in preserved_norm:
            logger.info(f"[FirstPagePreserver] Dropping generated section '{title}' -- already present in preserved template region.")
            continue
        kept.append(s)
    return kept


def build_document_with_preserved_first_page(
    template_path: str,
    sections: List[Dict[str, Any]],
    brand_cfg: Dict[str, Any],
    output_docx_path: str,
) -> str:
    """
    Build the final proposal .docx by:
      1. Opening a TEMPORARY COPY of the uploaded template (never the
         original upload path).
      2. Leaving everything up to the first explicit page break untouched,
         except patching any stale dates found there to today's date in the
         same format.
      3. Deleting everything after that page break (the template's own
         placeholder body, if any) and appending the AI-generated proposal
         sections in its place, styled using the template's own extracted
         brand config.

    Returns the path to the generated .docx.
    """
    from docx import Document
    import importlib
    try:
        from backend.scripts import proposal_generator as pg
    except ImportError:
        pg = importlib.import_module("scripts.proposal_generator")

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_copy = Path(tmp_dir) / "template_copy.docx"
        shutil.copyfile(template_path, tmp_copy)

        doc = Document(str(tmp_copy))
        split_idx = find_first_page_split_index(doc)

        # Belt-and-suspenders: even though document_generator.py already told
        # the section-writer prompt which headings exist in the template, if
        # the model wrote one anyway, drop it here before it's ever added to
        # the output document rather than shipping a visible duplicate.
        preserved_headings = detect_preserved_headings(doc, split_idx)
        sections = dedupe_sections_against_preserved(sections, preserved_headings)

        changed = _update_dates_in_preserved_region(doc, split_idx)
        if changed:
            logger.info(f"[FirstPagePreserver] Updated {changed} stale date field(s) on the preserved first page.")

        body = doc.element.body
        if split_idx is not None:
            # Remove everything AFTER the page-break paragraph, but keep the
            # trailing <w:sectPr> (page size/margins) if present so the
            # section layout for appended content stays intact.
            children = list(body)
            sect_pr = body.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr")
            for pos in range(len(children) - 1, split_idx, -1):
                el = children[pos]
                if el is sect_pr:
                    continue
                body.remove(el)
        else:
            # No explicit page break found -- the whole template is page 1.
            # Start the generated content on a fresh page.
            doc.add_page_break()

        _ensure_required_styles(doc)

        cfg = {"brand": brand_cfg, "sections": sections}
        for i, section in enumerate(sections, start=1):
            if i == 1:
                # We already ensured a fresh page above (or the template's own
                # page break); don't add a second one immediately.
                section = dict(section)
                section["page_break_before"] = False
            pg.add_section(doc, cfg, section, index=i)

        out_path = Path(output_docx_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))

    logger.info(
        f"[FirstPagePreserver] Generated document preserving template's first page -> {out_path} "
        f"(page break {'found' if split_idx is not None else 'not found, whole template kept as page 1'})"
    )
    return str(out_path)
