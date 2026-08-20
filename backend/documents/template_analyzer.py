"""
documents/template_analyzer.py
-------------------------------
Extracts branding + structured company content from user-uploaded .docx
templates (capability statements, dossiers, letterheads, etc.).

The template is used as a branding AND content reference — logos, colors,
fonts, header/footer design, PLUS tables, certifications, product suites,
and identifiers. A new document is always generated from scratch using
these extracted assets so templates are never mangled by content injection.

Designed to be layout-agnostic: works whether the file uses tables,
numbered sections, bold labels, or free-form paragraphs.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Identifier patterns (layout-agnostic; match labels in any order)
# ---------------------------------------------------------------------------
_ID_PATTERNS: List[Tuple[str, re.Pattern]] = [
    ("sap_anid", re.compile(r"\bANID[:\s]*([A-Z0-9]{8,})\b", re.I)),
    ("sap_anid", re.compile(r"\bSAP\s*(?:Business\s*)?Network\s*(?:ID|ANID)?[:\s]*([A-Z0-9]{8,})\b", re.I)),
    ("duns", re.compile(r"\bD[\-\s]?U[\-\s]?N[\-\s]?S[®\s]*(?:Number)?[:\s]*([0-9]{8,9})\b", re.I)),
    ("duns", re.compile(r"\b([0-9]{9})\b(?=[^\n]{0,40}D[\-\s]?U[\-\s]?N)", re.I)),
    ("sam_uei", re.compile(r"\b(?:SAM\.?gov\s*)?(?:Unique\s*Entity\s*ID|UEI)[:\s]*([A-Z0-9]{9,12})\b", re.I)),
    ("ncage", re.compile(r"\b(?:NCAGE|NATO\s*Supplier\s*Code|CAGE)[:\s]*([A-Z0-9]{4,6})\b", re.I)),
    ("ungm", re.compile(r"\bUNGM[:\s]*(?:Registration\s*ID)?[:\s]*([0-9]{5,})\b", re.I)),
    ("gem_seller", re.compile(r"\bGeM\s*(?:Seller\s*)?ID[:\s]*([A-Z0-9]{8,})\b", re.I)),
    ("msme_udyam", re.compile(r"\b(?:MSME\s*)?Udyam[:\s]*([A-Z0-9\-]{10,})\b", re.I)),
    ("dpiit", re.compile(r"\bDPIIT[:\s]*(?:DIPP)?[:\s]*([A-Z0-9]{6,})\b", re.I)),
    ("llpin", re.compile(r"\b(?:LLPIN|Corporate\s*Identity)[:\s]*([A-Z]{3}\-?[0-9]{4,})\b", re.I)),
    ("pan", re.compile(r"\bPAN[:\s]*([A-Z]{5}[0-9]{4}[A-Z])\b", re.I)),
    ("gstin", re.compile(r"\bGSTIN[:\s]*([0-9]{2}[A-Z]{5}[0-9]{4}[A-Z][A-Z0-9]Z[A-Z0-9])\b", re.I)),
    ("tan", re.compile(r"\bTAN[:\s]*([A-Z]{4}[0-9]{5}[A-Z])\b", re.I)),
    ("iec", re.compile(r"\bIEC[:\s]*([A-Z0-9]{8,12})\b", re.I)),
    ("iso_9001", re.compile(r"\bISO\s*9001(?::[0-9]{4})?\b[^\n]{0,80}", re.I)),
    ("iso_27001", re.compile(r"\bISO(?:/IEC)?\s*27001(?::[0-9]{4})?\b[^\n]{0,80}", re.I)),
    ("iso_45001", re.compile(r"\bISO\s*45001(?::[0-9]{4})?\b[^\n]{0,80}", re.I)),
    ("cmmi", re.compile(r"\bCMMI\s*(?:LEVEL\s*)?([1-5])\b[^\n]{0,60}", re.I)),
    ("soc2", re.compile(r"\bSOC\s*2\b[^\n]{0,60}", re.I)),
    ("gdpr", re.compile(r"\bGDPR\b[^\n]{0,60}", re.I)),
]


@dataclass
class BrandProfile:
    """
    Brand + content profile extracted from a user template or built from defaults.
    Drives document generation without modifying the original template.
    """
    # Fonts
    body_font: str = "Calibri"
    heading_font: str = "Calibri"

    # Colors (hex, no '#')
    accent_color: str = "1F3864"
    muted_color: str = "595959"
    heading_color: str = "1F3864"

    # Page setup
    page_width_inches: float = 8.5
    page_height_inches: float = 11.0
    left_margin_inches: float = 0.75
    right_margin_inches: float = 0.75
    top_margin_inches: float = 0.9
    bottom_margin_inches: float = 0.9

    # Logo (extracted to bytes so it can be embedded in new doc)
    logo_bytes: Optional[bytes] = None
    logo_width_inches: float = 1.2

    # Cover graphic
    cover_graphic_bytes: Optional[bytes] = None

    # Company info extracted from template
    company_name: str = "OrbitAvanya Tech LLP"
    website: str = ""
    address_line1: str = ""
    address_line2: str = ""
    phone: str = ""
    email: str = ""

    # Leadership names/titles found in the template
    leadership: List[Dict[str, str]] = field(default_factory=list)

    # Header/footer text patterns
    header_text: str = ""
    footer_text: str = ""

    # Raw text of the template's first page
    first_page_text: str = ""

    # ---- Dynamic structured content (layout-agnostic) ----
    # Full plain-text dump (paragraphs + tables) for AI context
    full_text: str = ""
    # Ordered section headings discovered in the doc
    sections: List[str] = field(default_factory=list)
    # Tables as list of {headers, rows, markdown}
    tables: List[Dict[str, Any]] = field(default_factory=list)
    # Key identifiers / certifications extracted via patterns
    identifiers: Dict[str, str] = field(default_factory=dict)
    # Bullet / product lines under major headings
    competencies: List[str] = field(default_factory=list)
    # Core product / solution suite lines
    products: List[str] = field(default_factory=list)

    # Source: 'template' | 'default' | 'default_fallback'
    source: str = "default"

    def to_brand_config_dict(self) -> Dict[str, Any]:
        """Convert to the brand config dict format used by proposal_generator."""
        out_dir = Path("downloads") / "extracted"
        out_dir.mkdir(parents=True, exist_ok=True)

        logo_path = ""
        cover_path = ""

        if self.logo_bytes:
            try:
                target_logo = out_dir / "logo.png"
                target_logo.write_bytes(self.logo_bytes)
                logo_path = str(target_logo.resolve())
            except Exception as e:
                logger.warning(f"Could not save extracted logo: {e}")

        if self.cover_graphic_bytes:
            try:
                target_cover = out_dir / "cover.png"
                target_cover.write_bytes(self.cover_graphic_bytes)
                cover_path = str(target_cover.resolve())
            except Exception as e:
                logger.warning(f"Could not save cover graphic: {e}")

        return {
            "company_name": self.company_name,
            "company_short": self.company_name.split()[0] if self.company_name else "Company",
            "logo_path": logo_path,
            "cover_graphic_path": cover_path,
            "body_font": self.body_font,
            "heading_font": self.heading_font,
            "accent_color": self.accent_color,
            "muted_color": self.muted_color,
            "address_line1": self.address_line1,
            "address_line2": self.address_line2,
            "phone": self.phone,
            "email": self.email,
            "website": self.website,
            "leadership": self.leadership,
            "first_page_text": self.first_page_text,
            "identifiers": self.identifiers,
            "sections": self.sections,
            "competencies": self.competencies,
            "products": self.products,
            "tables_markdown": [t.get("markdown", "") for t in self.tables if t.get("markdown")],
        }

    def to_company_profile_summary(self) -> str:
        """Plain-language summary of everything extracted from the template,
        meant for AI prompts so the model uses real contact / cert / product
        data instead of inventing or contradicting it."""
        lines: List[str] = []
        if self.company_name:
            lines.append(f"Company name: {self.company_name}")
        if self.website:
            lines.append(f"Website: {self.website}")
        if self.email:
            lines.append(f"Email: {self.email}")
        if self.phone:
            lines.append(f"Phone: {self.phone}")
        if self.address_line1 or self.address_line2:
            lines.append(
                f"Address: {', '.join(x for x in [self.address_line1, self.address_line2] if x)}"
            )
        if self.leadership:
            names = "; ".join(
                f"{p.get('name', '')} ({p.get('title', '')})"
                for p in self.leadership
                if p.get("name")
            )
            if names:
                lines.append(f"Leadership: {names}")

        if self.identifiers:
            lines.append("Identifiers & certifications:")
            for k, v in self.identifiers.items():
                lines.append(f"  - {k}: {v}")

        if self.sections:
            lines.append("Document sections: " + " | ".join(self.sections[:20]))

        if self.competencies:
            lines.append("Core competencies:")
            for c in self.competencies[:15]:
                lines.append(f"  - {c}")

        if self.products:
            lines.append("Products / solutions:")
            for p in self.products[:15]:
                lines.append(f"  - {p}")

        if self.tables:
            lines.append(f"Structured tables found: {len(self.tables)}")
            for i, t in enumerate(self.tables[:5], 1):
                md = (t.get("markdown") or "")[:1500]
                if md:
                    lines.append(f"--- Table {i} ---")
                    lines.append(md)

        # Cap overall size for prompt safety while keeping structure
        summary = "\n".join(l for l in lines if l)
        if len(summary) > 12000:
            summary = summary[:12000] + "\n…[truncated]"
        return summary


class TemplateAnalyzer:
    """
    Analyzes a user-uploaded .docx template to extract brand assets and
    structured company content. Does NOT modify the template.
    Layout-agnostic: tables, headings, free text, and mixed layouts all work.
    """

    def __init__(self, template_path: str | Path):
        self.template_path = Path(template_path)
        self._doc = None

    def _load_doc(self):
        if self._doc is None:
            from docx import Document
            self._doc = Document(str(self.template_path))
        return self._doc

    def analyze(self) -> BrandProfile:
        profile = BrandProfile(source="template")

        if not self.template_path.exists():
            logger.warning(f"Template not found: {self.template_path}. Using defaults.")
            profile.source = "default"
            return profile

        try:
            doc = self._load_doc()

            self._extract_fonts(doc, profile)
            self._extract_colors(doc, profile)
            self._extract_margins(doc, profile)
            self._extract_logo(doc, profile)
            self._extract_header_footer(doc, profile)
            self._extract_tables(doc, profile)
            self._extract_structure(doc, profile)
            self._extract_contact_details(doc, profile)
            self._extract_leadership(doc, profile)
            self._extract_identifiers(doc, profile)
            self._extract_company_name(doc, profile)
            self._extract_first_page_text(doc, profile)
            self._build_full_text(doc, profile)

            logger.info(
                f"[TemplateAnalyzer] Extracted brand+content from '{self.template_path.name}': "
                f"font={profile.body_font}, accent=#{profile.accent_color}, "
                f"tables={len(profile.tables)}, ids={len(profile.identifiers)}, "
                f"sections={len(profile.sections)}"
            )
        except Exception as e:
            logger.error(f"[TemplateAnalyzer] Failed to analyze template: {e}", exc_info=True)
            profile.source = "default_fallback"

        return profile

    # ------------------------------------------------------------------ fonts
    def _extract_fonts(self, doc, profile: BrandProfile):
        try:
            font_counts: Dict[str, int] = {}
            for para in doc.paragraphs[:80]:
                for run in para.runs:
                    fn = run.font.name
                    if fn and not fn.startswith("+"):
                        font_counts[fn] = font_counts.get(fn, 0) + 1
            # Also sample table cells
            for table in doc.tables[:5]:
                for row in table.rows[:10]:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                fn = run.font.name
                                if fn and not fn.startswith("+"):
                                    font_counts[fn] = font_counts.get(fn, 0) + 1

            if font_counts:
                profile.body_font = max(font_counts, key=lambda k: font_counts[k])

            for para in doc.paragraphs:
                if para.style and "heading" in (para.style.name or "").lower():
                    for run in para.runs:
                        if run.font.name and not run.font.name.startswith("+"):
                            profile.heading_font = run.font.name
                            break
                    break

            if not profile.heading_font or profile.heading_font == "Calibri":
                profile.heading_font = profile.body_font
        except Exception as e:
            logger.debug(f"Font extraction error: {e}")

    # ----------------------------------------------------------------- colors
    def _extract_colors(self, doc, profile: BrandProfile):
        try:
            for para in doc.paragraphs[:120]:
                style_name = (para.style.name or "").lower() if para.style else ""
                is_heading = "heading" in style_name or (para.runs and all(r.bold for r in para.runs if r.text.strip()))
                if not is_heading:
                    continue
                for run in para.runs:
                    if run.font.color and run.font.color.type is not None:
                        try:
                            rgb = run.font.color.rgb
                            if rgb:
                                hex_color = str(rgb)
                                if hex_color.upper() not in ("FFFFFF", "000000", "AUTO"):
                                    profile.accent_color = hex_color
                                    profile.heading_color = hex_color
                                    return
                        except Exception:
                            pass
        except Exception as e:
            logger.debug(f"Color extraction error: {e}")

    # ---------------------------------------------------------------- margins
    def _extract_margins(self, doc, profile: BrandProfile):
        try:
            section = doc.sections[0]
            emu_per_inch = 914400
            if section.left_margin:
                profile.left_margin_inches = round(section.left_margin / emu_per_inch, 2)
            if section.right_margin:
                profile.right_margin_inches = round(section.right_margin / emu_per_inch, 2)
            if section.top_margin:
                profile.top_margin_inches = round(section.top_margin / emu_per_inch, 2)
            if section.bottom_margin:
                profile.bottom_margin_inches = round(section.bottom_margin / emu_per_inch, 2)
            if section.page_width:
                profile.page_width_inches = round(section.page_width / emu_per_inch, 2)
            if section.page_height:
                profile.page_height_inches = round(section.page_height / emu_per_inch, 2)
        except Exception as e:
            logger.debug(f"Margin extraction error: {e}")

    # ------------------------------------------------------------------- logo
    def _extract_logo(self, doc, profile: BrandProfile):
        """Prefer header image; fall back to first body image."""
        try:
            ns_a = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
            ns_r = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"

            def _from_paragraphs(paragraphs, part_related):
                for para in paragraphs:
                    for run in para.runs:
                        for shape in run._r.findall(f".//{ns_a}blip"):
                            embed = shape.get(f"{ns_r}embed")
                            if not embed:
                                continue
                            img_part = part_related.get(embed)
                            if img_part and getattr(img_part, "blob", None):
                                return img_part.blob
                return None

            # 1) Headers
            for section in doc.sections:
                for header in (section.header, section.first_page_header):
                    if header is None:
                        continue
                    blob = _from_paragraphs(header.paragraphs, doc.part.related_parts)
                    if blob:
                        profile.logo_bytes = blob
                        logger.info("[TemplateAnalyzer] Extracted logo from header.")
                        return

            # 2) First body image (common in cover pages without header logo)
            blob = _from_paragraphs(doc.paragraphs[:30], doc.part.related_parts)
            if blob:
                profile.logo_bytes = blob
                logger.info("[TemplateAnalyzer] Extracted logo from body.")
        except Exception as e:
            logger.debug(f"Logo extraction error: {e}")

    # ---------------------------------------------------------- header/footer
    def _extract_header_footer(self, doc, profile: BrandProfile):
        try:
            section = doc.sections[0]
            if section.header:
                texts = [p.text.strip() for p in section.header.paragraphs if p.text.strip()]
                profile.header_text = " | ".join(texts)
            if section.footer:
                texts = [p.text.strip() for p in section.footer.paragraphs if p.text.strip()]
                profile.footer_text = " | ".join(texts)
                for text in texts:
                    if "www." in text.lower() or ".com" in text.lower():
                        profile.website = text.strip()
        except Exception as e:
            logger.debug(f"Header/footer extraction error: {e}")

    # ----------------------------------------------------------------- tables
    def _extract_tables(self, doc, profile: BrandProfile):
        """Extract every table as headers + rows + markdown. Critical for
        accreditation matrices and similar capability-statement layouts."""
        try:
            tables_out: List[Dict[str, Any]] = []
            for table in doc.tables:
                rows_raw: List[List[str]] = []
                for row in table.rows:
                    cells = []
                    for cell in row.cells:
                        # Join multi-paragraph cells; collapse whitespace
                        cell_text = " ".join(
                            p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()
                        )
                        cells.append(re.sub(r"\s+", " ", cell_text).strip())
                    # Skip fully empty rows
                    if any(c for c in cells):
                        rows_raw.append(cells)

                if not rows_raw:
                    continue

                # Normalize column count
                max_cols = max(len(r) for r in rows_raw)
                rows_raw = [r + [""] * (max_cols - len(r)) for r in rows_raw]

                headers = rows_raw[0]
                body = rows_raw[1:] if len(rows_raw) > 1 else []

                # Build markdown table
                def _esc(s: str) -> str:
                    return s.replace("|", "\\|")

                md_lines = [
                    "| " + " | ".join(_esc(h) for h in headers) + " |",
                    "| " + " | ".join("---" for _ in headers) + " |",
                ]
                for r in body:
                    md_lines.append("| " + " | ".join(_esc(c) for c in r) + " |")

                tables_out.append({
                    "headers": headers,
                    "rows": body,
                    "markdown": "\n".join(md_lines),
                    "row_count": len(body),
                    "col_count": max_cols,
                })

            profile.tables = tables_out
        except Exception as e:
            logger.debug(f"Table extraction error: {e}")

    # -------------------------------------------------------------- structure
    def _extract_structure(self, doc, profile: BrandProfile):
        """Discover section headings and collect bullets under product /
        competency style sections without hardcoding fixed templates."""
        try:
            sections: List[str] = []
            competencies: List[str] = []
            products: List[str] = []
            current_bucket: Optional[str] = None

            product_keys = (
                "product", "solution", "suite", "offering", "service",
                "core enterprise", "platform",
            )
            competency_keys = (
                "competenc", "core compet", "capability", "strength",
                "expertise", "speciali",
            )

            for para in doc.paragraphs:
                text = (para.text or "").strip()
                if not text:
                    continue

                style_name = (para.style.name or "").lower() if para.style else ""
                is_heading = (
                    "heading" in style_name
                    or bool(re.match(r"^\d+[\.\)]\s+\S", text))
                    or (len(text) < 120 and text.isupper() and len(text.split()) <= 12)
                    or (
                        para.runs
                        and all(r.bold for r in para.runs if r.text.strip())
                        and len(text) < 140
                        and not text.endswith(".")
                    )
                )

                if is_heading:
                    # Strip leading numbering for cleaner section list
                    clean = re.sub(r"^\d+[\.\)]\s*", "", text).strip()
                    if clean and clean not in sections:
                        sections.append(clean)
                    lower = clean.lower()
                    if any(k in lower for k in product_keys):
                        current_bucket = "products"
                    elif any(k in lower for k in competency_keys):
                        current_bucket = "competencies"
                    else:
                        current_bucket = None
                    continue

                # Bullets / short lines under active bucket
                is_bullet = text.startswith(("•", "-", "–", "—", "*", "·")) or (
                    para.style and "list" in style_name
                )
                cleaned = re.sub(r"^[•\-–—*·]\s*", "", text).strip()
                if current_bucket == "products" and cleaned and len(cleaned) < 300:
                    if is_bullet or len(cleaned) < 160:
                        products.append(cleaned)
                elif current_bucket == "competencies" and cleaned and len(cleaned) < 300:
                    if is_bullet or len(cleaned) < 160:
                        competencies.append(cleaned)

            profile.sections = sections[:40]
            profile.products = list(dict.fromkeys(products))[:30]
            profile.competencies = list(dict.fromkeys(competencies))[:30]
        except Exception as e:
            logger.debug(f"Structure extraction error: {e}")

    # ------------------------------------------------------- contact details
    _EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
    _PHONE_RE = re.compile(r"(\+?\d[\d\-\.\(\) ]{7,}\d)")
    _WEBSITE_RE = re.compile(
        r"(?:https?://)?(?:www\.)?[A-Za-z0-9\-]+(?:\.[A-Za-z0-9\-]+)+(?:/\S*)?"
    )
    _LEADERSHIP_TITLE_RE = re.compile(
        r"\b(Chief Executive Officer|CEO|Chief Operating Officer|COO|"
        r"Chief Technology Officer|CTO|Chief Financial Officer|CFO|"
        r"President|Founder|Co-Founder|Managing Director|Director|"
        r"Vice President|VP|Owner|Principal|Partner|Authorized Signatory)\b",
        re.IGNORECASE,
    )
    _ADDRESS_RE = re.compile(
        r"(?:\d{1,4}[/\-]?\d*[,\s]+[A-Za-z][^\n]{10,120}"
        r"(?:Mumbai|Delhi|Bangalore|Hyderabad|Pune|Chennai|London|Dubai|"
        r"Riyadh|Houston|New York|California|Maharashtra|India|UK|USA|"
        r"United Kingdom|United States)[^\n]{0,40})",
        re.I,
    )

    def _all_text_sources(self, doc) -> List[str]:
        """Paragraphs + table cells + headers/footers — full document surface."""
        texts: List[str] = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                texts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = " ".join(
                        p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()
                    )
                    if cell_text:
                        texts.append(cell_text)
        for section in doc.sections:
            for part in (
                section.header,
                section.footer,
                section.first_page_header,
                section.first_page_footer,
            ):
                if part is None:
                    continue
                for p in part.paragraphs:
                    if p.text and p.text.strip():
                        texts.append(p.text)
        return texts

    def _extract_contact_details(self, doc, profile: BrandProfile):
        try:
            all_text = "\n".join(self._all_text_sources(doc))

            email_match = self._EMAIL_RE.search(all_text)
            if email_match:
                profile.email = email_match.group(0)

            if not profile.website:
                for candidate in self._WEBSITE_RE.finditer(all_text):
                    val = candidate.group(0).rstrip(".,);]")
                    if profile.email and val in profile.email:
                        continue
                    lower = val.lower()
                    if any(
                        lower.endswith(ext) or f".{ext}/" in lower
                        for ext in ("com", "org", "net", "io", "co", "in", "uk", "ae", "sa")
                    ) or "www." in lower:
                        # Prefer full https URL if present
                        if not val.startswith("http"):
                            profile.website = "https://" + val.lstrip("/")
                        else:
                            profile.website = val
                        break

            if not profile.phone:
                for phone_match in self._PHONE_RE.finditer(all_text):
                    candidate = phone_match.group(1).strip()
                    digit_count = sum(c.isdigit() for c in candidate)
                    if 7 <= digit_count <= 15:
                        profile.phone = candidate
                        break

            # Address heuristic
            addr = self._ADDRESS_RE.search(all_text)
            if addr:
                full = re.sub(r"\s+", " ", addr.group(0)).strip()
                if len(full) > 20:
                    # Split roughly in half for line1/line2 if long
                    if len(full) > 70 and "," in full:
                        parts = full.split(",")
                        mid = len(parts) // 2
                        profile.address_line1 = ", ".join(parts[:mid]).strip()
                        profile.address_line2 = ", ".join(parts[mid:]).strip()
                    else:
                        profile.address_line1 = full
        except Exception as e:
            logger.debug(f"Contact detail extraction error: {e}")

    # ------------------------------------------------------------ leadership
    def _extract_leadership(self, doc, profile: BrandProfile):
        try:
            leadership: List[Dict[str, str]] = []
            seen = set()
            for text in self._all_text_sources(doc):
                title_match = self._LEADERSHIP_TITLE_RE.search(text)
                if not title_match:
                    continue
                title = title_match.group(0)
                name = ""
                for sep in ("—", "–", "-", ",", "(", ":"):
                    if sep not in text:
                        continue
                    left, _, right = text.partition(sep)
                    for candidate_raw in (left, right):
                        candidate = candidate_raw.strip(" ()").strip()
                        if title.lower() in candidate.lower():
                            continue
                        words = candidate.split()
                        if (
                            1 < len(words) <= 5
                            and all(w[:1].isupper() for w in words if w and w[0].isalpha())
                            and not any(c.isdigit() for c in candidate)
                        ):
                            name = candidate
                            break
                    if name:
                        break

                # "Signatory: Ranjeet Kumar (Founder & CEO)" style
                if not name:
                    m = re.search(
                        r"(?:Signatory|Name)[:\s]+([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){1,3})",
                        text,
                    )
                    if m:
                        name = m.group(1).strip()

                if name and (name, title) not in seen:
                    seen.add((name, title))
                    leadership.append({"name": name, "title": title})

            profile.leadership = leadership[:8]
        except Exception as e:
            logger.debug(f"Leadership extraction error: {e}")

    # ---------------------------------------------------------- identifiers
    def _extract_identifiers(self, doc, profile: BrandProfile):
        """Pull certifications and registry IDs from any layout (tables or prose)."""
        try:
            all_text = "\n".join(self._all_text_sources(doc))
            found: Dict[str, str] = {}
            for key, pattern in _ID_PATTERNS:
                if key in found:
                    continue
                m = pattern.search(all_text)
                if not m:
                    continue
                value = (m.group(1) if m.lastindex else m.group(0)).strip()
                value = re.sub(r"\s+", " ", value)
                if value:
                    found[key] = value
            profile.identifiers = found
        except Exception as e:
            logger.debug(f"Identifier extraction error: {e}")

    # --------------------------------------------------------- company name
    def _extract_company_name(self, doc, profile: BrandProfile):
        """Prefer explicit legal-entity labels; fall back to first strong heading."""
        try:
            all_text = "\n".join(self._all_text_sources(doc))
            m = re.search(
                r"(?:Legal\s*Entity\s*Name|Company\s*Name|Organisation\s*Name)[:\s|]*"
                r"([A-Z][A-Za-z0-9 &.\-]{3,80}(?:LLP|Ltd|Limited|PLC|Inc|LLC|Pvt\.?\s*Ltd)?)",
                all_text,
                re.I,
            )
            if m:
                profile.company_name = m.group(1).strip(" |*")
                return

            # First non-empty heading-like paragraph that looks like a company name
            for para in doc.paragraphs[:15]:
                text = (para.text or "").strip()
                if not text or len(text) > 90:
                    continue
                style = (para.style.name or "").lower() if para.style else ""
                if "heading" in style or (para.runs and any(r.bold for r in para.runs)):
                    if any(
                        tok in text.upper()
                        for tok in ("LLP", "LTD", "LIMITED", "INC", "LLC", "PLC", "PVT")
                    ):
                        profile.company_name = text
                        return
        except Exception as e:
            logger.debug(f"Company name extraction error: {e}")

    # -------------------------------------------------------- first page text
    def _extract_first_page_text(self, doc, profile: BrandProfile):
        try:
            from documents.bidforge.first_page_preserver import iter_first_page_paragraph_texts
            texts = list(iter_first_page_paragraph_texts(doc))
            # Append first table markdown so accreditation matrix on page 1 is kept
            table_bits = []
            for t in profile.tables[:2]:
                if t.get("markdown"):
                    table_bits.append(t["markdown"])
            combined = "\n".join(t for t in texts if t.strip())
            if table_bits:
                combined = combined + "\n\n" + "\n\n".join(table_bits)
            profile.first_page_text = combined
        except Exception as e:
            logger.debug(f"First-page text extraction error: {e}")
            # Fallback: first N paragraphs + first table
            try:
                paras = [p.text for p in doc.paragraphs[:40] if p.text and p.text.strip()]
                profile.first_page_text = "\n".join(paras)
                if profile.tables:
                    profile.first_page_text += "\n\n" + profile.tables[0].get("markdown", "")
            except Exception:
                pass

    # ----------------------------------------------------------- full text
    def _build_full_text(self, doc, profile: BrandProfile):
        """Complete document surface for downstream AI context (capped)."""
        try:
            parts = self._all_text_sources(doc)
            for t in profile.tables:
                if t.get("markdown"):
                    parts.append(t["markdown"])
            full = "\n".join(parts)
            profile.full_text = full[:20000] if len(full) > 20000 else full
        except Exception as e:
            logger.debug(f"Full text build error: {e}")


def analyze_template(template_path: str | Path) -> BrandProfile:
    """Convenience function to analyze a template and return its BrandProfile."""
    return TemplateAnalyzer(template_path).analyze()


def get_default_brand_profile() -> BrandProfile:
    """Return the default OrbitAvanya brand profile."""
    from documents.brand_config import get_brand_config

    cfg = get_brand_config()

    profile = BrandProfile(source="default")
    profile.body_font = cfg.get("body_font", "Calibri")
    profile.heading_font = cfg.get("heading_font", "Calibri")
    profile.accent_color = cfg.get("accent_color", "1F3864")
    profile.muted_color = cfg.get("muted_color", "595959")
    profile.company_name = cfg.get("company_name", "OrbitAvanya Tech LLP")
    profile.website = cfg.get("website", "")
    profile.address_line1 = cfg.get("address_line1", "")
    profile.address_line2 = cfg.get("address_line2", "")
    profile.phone = cfg.get("phone", "")
    profile.email = cfg.get("email", "")
    profile.identifiers = cfg.get("identifiers") or {}
    profile.competencies = cfg.get("competencies") or []
    profile.products = cfg.get("products") or []

    logo_path = cfg.get("logo_path", "")
    if logo_path and Path(logo_path).exists():
        try:
            profile.logo_bytes = Path(logo_path).read_bytes()
        except Exception:
            pass

    cover_path = cfg.get("cover_graphic_path", "")
    if cover_path and Path(cover_path).exists():
        try:
            profile.cover_graphic_bytes = Path(cover_path).read_bytes()
        except Exception:
            pass

    return profile
