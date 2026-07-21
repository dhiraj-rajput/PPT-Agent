"""
documents/template_analyzer.py
-------------------------------
Extracts branding assets from user-uploaded .docx templates.
The template is used ONLY as a branding reference — logos, colors,
fonts, header/footer design. A new document is always generated from scratch
using these extracted brand assets.

This prevents the 'fill-in-the-blank' problem where templates get
mangled by content injection.
"""

from __future__ import annotations

import logging
import re
import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


@dataclass
class BrandProfile:
    """
    Brand profile extracted from a user template or built from defaults.
    Drives document generation without modifying the original template.
    """
    # Fonts
    body_font: str = "Calibri"
    heading_font: str = "Calibri"

    # Colors (hex, no '#')
    accent_color: str = "1F3864"
    muted_color: str = "595959"
    heading_color: str = "1F3864"

    # Page setup
    page_width_inches: float = 8.5
    page_height_inches: float = 11.0
    left_margin_inches: float = 0.75
    right_margin_inches: float = 0.75
    top_margin_inches: float = 0.9
    bottom_margin_inches: float = 0.9

    # Logo (extracted to bytes so it can be embedded in new doc)
    logo_bytes: Optional[bytes] = None
    logo_width_inches: float = 1.2

    # Cover graphic
    cover_graphic_bytes: Optional[bytes] = None

    # Company info extracted from template
    company_name: str = "OrbitAvanya Tech LLP"
    website: str = ""
    address_line1: str = ""
    address_line2: str = ""
    phone: str = ""

    # Header/footer text patterns
    header_text: str = ""
    footer_text: str = ""

    # Source: 'template' | 'default'
    source: str = "default"

    def to_brand_config_dict(self) -> Dict[str, Any]:
        """Convert to the brand config dict format used by proposal_generator."""
        import tempfile
        import os

        logo_path = ""
        cover_path = ""

        # Save logo bytes to temp file if available
        if self.logo_bytes:
            try:
                suffix = ".png"
                with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
                    f.write(self.logo_bytes)
                    logo_path = f.name
            except Exception as e:
                logger.warning(f"Could not save extracted logo: {e}")

        if self.cover_graphic_bytes:
            try:
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
                    f.write(self.cover_graphic_bytes)
                    cover_path = f.name
            except Exception as e:
                logger.warning(f"Could not save cover graphic: {e}")

        return {
            "company_name": self.company_name,
            "company_short": self.company_name.split()[0] if self.company_name else "Company",
            "logo_path": logo_path,
            "cover_graphic_path": cover_path,
            "body_font": self.body_font,
            "heading_font": self.heading_font,
            "accent_color": self.accent_color,
            "muted_color": self.muted_color,
            "address_line1": self.address_line1,
            "address_line2": self.address_line2,
            "phone": self.phone,
            "website": self.website,
        }


class TemplateAnalyzer:
    """
    Analyzes a user-uploaded .docx template to extract brand assets.
    Does NOT modify the template.
    """

    def __init__(self, template_path: str | Path):
        self.template_path = Path(template_path)
        self._doc = None

    def _load_doc(self):
        """Lazy-load the template document."""
        if self._doc is None:
            from docx import Document
            self._doc = Document(str(self.template_path))
        return self._doc

    def analyze(self) -> BrandProfile:
        """
        Extract all branding information from the template.
        Returns a BrandProfile with all available brand data.
        """
        profile = BrandProfile(source="template")

        if not self.template_path.exists():
            logger.warning(f"Template not found: {self.template_path}. Using defaults.")
            profile.source = "default"
            return profile

        try:
            doc = self._load_doc()

            self._extract_fonts(doc, profile)
            self._extract_colors(doc, profile)
            self._extract_margins(doc, profile)
            self._extract_logo(doc, profile)
            self._extract_header_footer(doc, profile)

            logger.info(
                f"[TemplateAnalyzer] Extracted brand from '{self.template_path.name}': "
                f"font={profile.body_font}, accent=#{profile.accent_color}"
            )
        except Exception as e:
            logger.error(f"[TemplateAnalyzer] Failed to analyze template: {e}")
            profile.source = "default_fallback"

        return profile

    def _extract_fonts(self, doc, profile: BrandProfile):
        """Extract primary font from the template's paragraph styles."""
        try:
            font_counts: Dict[str, int] = {}
            for para in doc.paragraphs[:50]:  # Sample first 50 paragraphs
                for run in para.runs:
                    fn = run.font.name
                    if fn and not fn.startswith("+"):
                        font_counts[fn] = font_counts.get(fn, 0) + 1

            if font_counts:
                # Most common font = body font
                body = max(font_counts, key=lambda k: font_counts[k])
                profile.body_font = body

            # Check heading styles for heading font
            for para in doc.paragraphs:
                if para.style and "heading" in para.style.name.lower():
                    for run in para.runs:
                        if run.font.name and not run.font.name.startswith("+"):
                            profile.heading_font = run.font.name
                            break
                    break

            if not profile.heading_font or profile.heading_font == "Calibri":
                profile.heading_font = profile.body_font
        except Exception as e:
            logger.debug(f"Font extraction error: {e}")

    def _extract_colors(self, doc, profile: BrandProfile):
        """Extract accent color from heading styles or bold text."""
        try:
            for para in doc.paragraphs[:100]:
                if para.style and "heading" in para.style.name.lower():
                    for run in para.runs:
                        if run.font.color and run.font.color.type is not None:
                            try:
                                rgb = run.font.color.rgb
                                if rgb:
                                    hex_color = str(rgb)
                                    # Avoid very light or white colors
                                    if hex_color not in ('FFFFFF', 'ffffff', '000000'):
                                        profile.accent_color = hex_color
                                        profile.heading_color = hex_color
                                        return
                            except Exception:
                                pass
        except Exception as e:
            logger.debug(f"Color extraction error: {e}")

    def _extract_margins(self, doc, profile: BrandProfile):
        """Extract page margins from the template's section settings."""
        try:
            from docx.shared import Inches
            section = doc.sections[0]
            emu_per_inch = 914400

            if section.left_margin:
                profile.left_margin_inches = round(section.left_margin / emu_per_inch, 2)
            if section.right_margin:
                profile.right_margin_inches = round(section.right_margin / emu_per_inch, 2)
            if section.top_margin:
                profile.top_margin_inches = round(section.top_margin / emu_per_inch, 2)
            if section.bottom_margin:
                profile.bottom_margin_inches = round(section.bottom_margin / emu_per_inch, 2)
            if section.page_width:
                profile.page_width_inches = round(section.page_width / emu_per_inch, 2)
            if section.page_height:
                profile.page_height_inches = round(section.page_height / emu_per_inch, 2)
        except Exception as e:
            logger.debug(f"Margin extraction error: {e}")

    def _extract_logo(self, doc, profile: BrandProfile):
        """Extract logo image from the template header."""
        try:
            for section in doc.sections:
                for header in (section.header, section.first_page_header):
                    if header is None:
                        continue
                    for para in header.paragraphs:
                        for run in para.runs:
                            for shape in run._r.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                                embed = shape.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if embed:
                                    try:
                                        img_part = doc.part.related_parts.get(embed)
                                        if img_part:
                                            profile.logo_bytes = img_part.blob
                                            logger.info("[TemplateAnalyzer] Extracted logo from header.")
                                            return
                                    except Exception:
                                        pass
        except Exception as e:
            logger.debug(f"Logo extraction error: {e}")

    def _extract_header_footer(self, doc, profile: BrandProfile):
        """Extract text content from template header and footer."""
        try:
            section = doc.sections[0]
            if section.header:
                texts = [p.text.strip() for p in section.header.paragraphs if p.text.strip()]
                profile.header_text = " | ".join(texts)
            if section.footer:
                texts = [p.text.strip() for p in section.footer.paragraphs if p.text.strip()]
                profile.footer_text = " | ".join(texts)
                # Try to extract company name, website from footer
                for text in texts:
                    if 'www.' in text or '.com' in text:
                        profile.website = text.strip()
        except Exception as e:
            logger.debug(f"Header/footer extraction error: {e}")


def analyze_template(template_path: str | Path) -> BrandProfile:
    """Convenience function to analyze a template and return its BrandProfile."""
    return TemplateAnalyzer(template_path).analyze()


def get_default_brand_profile() -> BrandProfile:
    """Return the default OrbitAvanya brand profile."""
    from documents.brand_config import get_brand_config, PROJECT_ROOT
    cfg = get_brand_config()

    profile = BrandProfile(source="default")
    profile.body_font = cfg.get("body_font", "Calibri")
    profile.heading_font = cfg.get("heading_font", "Calibri")
    profile.accent_color = cfg.get("accent_color", "1F3864")
    profile.muted_color = cfg.get("muted_color", "595959")
    profile.company_name = cfg.get("company_name", "OrbitAvanya Tech LLP")
    profile.website = cfg.get("website", "")
    profile.address_line1 = cfg.get("address_line1", "")
    profile.address_line2 = cfg.get("address_line2", "")
    profile.phone = cfg.get("phone", "")

    # Load logo bytes from disk
    logo_path = cfg.get("logo_path", "")
    if logo_path and Path(logo_path).exists():
        try:
            profile.logo_bytes = Path(logo_path).read_bytes()
        except Exception:
            pass

    cover_path = cfg.get("cover_graphic_path", "")
    if cover_path and Path(cover_path).exists():
        try:
            profile.cover_graphic_bytes = Path(cover_path).read_bytes()
        except Exception:
            pass

    return profile
