"""
bidforge/template_filler.py
----------------------------
Populates a user-uploaded .docx template with AI-generated proposal content,
instead of building the document from scratch.

Design (deliberately NOT tied to a specific template-authoring convention —
works with any normal Word document that uses heading styles):

  1. Walk the template's paragraphs and find heading paragraphs (Word's
     "Heading 1"/"Heading 2"/"Title" styles).
  2. Fuzzy-match each heading's text against known proposal section keywords
     (Executive Summary, Scope, Pricing, Timeline, Terms, ...).
  3. Insert the generated content for a matched section directly after that
     heading, using the *template's own* body-text style — so the
     template's fonts/branding/spacing are preserved exactly.
  4. Any generated section that doesn't match an existing heading in the
     template is appended at the end (after a page break) under its own
     heading, so nothing is silently dropped.

Because we only ever touch body paragraphs — never section/header/footer
XML — the template's headers, footers, logos, and page numbering carry
through automatically and the document is free to run to however many
pages the content needs (no page cap is imposed anywhere in this code).
"""

from __future__ import annotations

from typing import Any, Dict, List, Optional

from docx import Document as OpenDocument
from docx.document import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.helpers import setup_logger

logger = setup_logger(__name__)


# Section key -> heading keywords used to match it against the template's own headings
_SECTION_KEYWORDS = {
    "executive_summary": ["executive summary", "overview", "introduction"],
    "scope_of_work": ["scope of work", "scope", "solution", "approach", "products and services"],
    "pricing_table": ["pricing", "cost", "investment", "budget", "fees"],
    "competitive_positioning": ["competitive", "value proposition", "why us", "differentiat"],
    "timeline": ["timeline", "schedule", "implementation", "milestones", "roadmap"],
    "terms": ["terms", "conditions", "sla", "warranty", "support"],
    "next_steps": ["next steps", "conclusion", "contact"],
}

_SECTION_TITLES = {
    "executive_summary": "Executive Summary",
    "scope_of_work": "Scope of Work",
    "pricing_table": "Pricing",
    "competitive_positioning": "Competitive Positioning",
    "timeline": "Implementation Timeline",
    "terms": "Terms & Conditions",
    "next_steps": "Next Steps",
}


def fill_template(template_path: str, sections: Dict[str, Any], output_path: str) -> str:
    """
    sections: {
      "executive_summary": "text...",
      "scope_of_work": ["bullet", "bullet", ...] or "text",
      "pricing_table": {"headers": [...], "rows": [[...], ...]},
      "competitive_positioning": "text",
      "timeline": [{"phase": ..., "duration": ..., "focus": ...}, ...] or "text",
      "terms": ["term1", "term2", ...],
      "next_steps": "text",
    }
    """
    doc = OpenDocument(template_path)
    body_style = _guess_body_style(doc)

    headings = _find_headings(doc)
    used_section_keys = set()

    for heading_para, heading_text in headings:
        section_key = _match_section(heading_text)
        if section_key and section_key not in used_section_keys:
            anchor = heading_para
            anchor = _insert_section_content(doc, anchor, sections.get(section_key), body_style)
            used_section_keys.add(section_key)

    # Anything not matched to an existing heading gets appended at the end.
    remaining = [k for k in sections.keys() if k in _SECTION_TITLES and k not in used_section_keys and sections.get(k)]
    if remaining:
        doc.add_page_break()
        for key in remaining:
            heading_para = doc.add_heading(_SECTION_TITLES[key], level=1)
            _insert_section_content(doc, heading_para, sections.get(key), body_style, append_mode=True)

    doc.save(output_path)
    logger.info(
        f"[BidForge:TemplateFiller] Filled template. Matched headings: {sorted(used_section_keys)}. "
        f"Appended (no matching heading): {remaining}"
    )
    return output_path


def _guess_body_style(doc: Document) -> Optional[str]:
    for p in doc.paragraphs:
        if p.style and p.style.name and not p.style.name.lower().startswith(("heading", "title")) and p.text.strip():
            return p.style.name
    return "Normal"


def _find_headings(doc: Document) -> List[tuple]:
    headings = []
    for p in doc.paragraphs:
        style_name = (p.style.name if p.style else "") or ""
        if style_name.lower().startswith("heading") or style_name.lower() == "title":
            if p.text.strip():
                headings.append((p, p.text.strip()))
    return headings


def _match_section(heading_text: str) -> Optional[str]:
    text_lower = heading_text.lower()
    for key, keywords in _SECTION_KEYWORDS.items():
        if any(kw in text_lower for kw in keywords):
            return key
    return None


def _insert_paragraph_after(paragraph: Paragraph, text: str = "", style: Optional[str] = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def _insert_table_after(paragraph: Paragraph, doc: Document, headers: List[str], rows: List[List[str]]) -> Paragraph:
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr_cells = table.rows[0].cells
    for cell, text in zip(hdr_cells, headers):
        cell.text = str(text)
    for row_data in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, row_data):
            cell.text = str(text)

    # Move the table's XML element to right after `paragraph`
    paragraph._p.addnext(table._tbl)
    # Return a fresh trailing paragraph so subsequent inserts continue below the table
    trailing = OxmlElement("w:p")
    table._tbl.addnext(trailing)
    return Paragraph(trailing, paragraph._parent)


def _insert_section_content(
    doc: Document,
    anchor: Paragraph,
    content: Any,
    body_style: Optional[str],
    append_mode: bool = False,
) -> Paragraph:
    """Inserts `content` (str, list of str, list of dicts, or a pricing-table dict)
    directly after `anchor`, returning the new last paragraph so callers can chain
    further inserts in document order."""
    if not content:
        return anchor

    if isinstance(content, dict) and "headers" in content and "rows" in content:
        return _insert_table_after(anchor, doc, content["headers"], content["rows"])

    if isinstance(content, str):
        for para_text in content.split("\n\n"):
            para_text = para_text.strip()
            if para_text:
                anchor = _insert_paragraph_after(anchor, para_text, style=body_style)
        return anchor

    if isinstance(content, list):
        for item in content:
            if isinstance(item, dict):
                # e.g. timeline phase dicts -> render as a compact line
                line = " — ".join(str(v) for v in item.values() if v)
                anchor = _insert_paragraph_after(anchor, line, style=body_style)
            else:
                p = _insert_paragraph_after(anchor, str(item), style=None)
                try:
                    p.style = "List Bullet"
                except KeyError:
                    pass
                anchor = p
        return anchor

    return anchor
