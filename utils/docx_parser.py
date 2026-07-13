"""
utils/docx_parser.py
---------------------
Extracts design assets and content structure from the WK T360 Proposal DOCX template.

Extracted assets:
  - Color palette from document theme
  - Font names and sizes
  - Section structure (TOC headings)
  - Header/footer layout
  - Page dimensions and margins
  - Phase table structure
  - Company profile data embedded in the document

Usage:
    from utils.docx_parser import DOCXTemplateParser
    assets = DOCXTemplateParser().parse()
    print(assets.colors.primary)
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Asset data classes
# ---------------------------------------------------------------------------

@dataclass
class DOCXColors:
    """Color palette extracted from the WK T360 DOCX theme."""
    primary:     str = "#44546A"    # Dark Navy — primary headings, header bg
    secondary:   str = "#ED7D31"    # Orange Accent — dividers, callouts
    accent_blue: str = "#5B9BD5"    # Blue Accent — table headers
    mid_blue:    str = "#4472C4"    # Mid Blue — alternate table header
    gold:        str = "#FFC000"    # Gold — key highlights
    light_grey:  str = "#E7E6E6"    # Light Grey — row alternates, bg
    cover_bg:    str = "#F2F2F2"    # Cover info block background
    conf_bar:    str = "#D9D9D9"    # Confidentiality bar background
    white:       str = "#FFFFFF"
    black:       str = "#000000"
    green:       str = "#70AD47"    # Success indicators


@dataclass
class DOCXFonts:
    """Font configuration from WK T360 template."""
    primary:     str = "Fira Sans Light"
    bold:        str = "Fira Sans SemiBold"
    # ReportLab fallbacks when Fira Sans is not embedded
    primary_rl:  str = "Helvetica"
    bold_rl:     str = "Helvetica-Bold"


@dataclass
class DOCXHeader:
    """Header table layout: two-column table."""
    left_label:  str = "Customer"
    right_label: str = "Proposal"
    left_value:  str = "Issuing Agency"     # replaced at render time
    right_value: str = "OrbitAvanya Tech"


@dataclass
class DOCXFooter:
    """Footer table layout: two-column table."""
    left_text:   str = ""                         # page number goes here
    right_text:  str = "OrbitAvanya Tech\nwww.orbitavanyatech.com"


@dataclass
class DOCXPageLayout:
    """Page dimensions matching the DOCX (US Letter)."""
    width_cm:        float = 21.59
    height_cm:       float = 27.94
    left_margin_cm:  float = 1.27
    right_margin_cm: float = 1.27
    top_margin_cm:   float = 1.5
    bottom_margin_cm: float = 1.5


@dataclass
class DOCXTypography:
    """Font sizes used in the template."""
    cover_title_pt:     float = 24.0
    section_heading_pt: float = 22.0
    sub_heading_pt:     float = 16.0
    body_pt:            float = 10.0
    table_header_pt:    float = 10.0
    table_body_pt:      float = 9.0
    caption_pt:         float = 8.0


@dataclass
class DOCXPhaseRow:
    """A single implementation phase row."""
    phase:     str = ""
    duration:  str = ""
    focus:     str = ""


@dataclass
class DOCXTemplateAssets:
    """All extracted design assets from WK T360 Proposal DOCX."""
    colors:      DOCXColors     = field(default_factory=DOCXColors)
    fonts:       DOCXFonts      = field(default_factory=DOCXFonts)
    header:      DOCXHeader     = field(default_factory=DOCXHeader)
    footer:      DOCXFooter     = field(default_factory=DOCXFooter)
    page_layout: DOCXPageLayout = field(default_factory=DOCXPageLayout)
    typography:  DOCXTypography = field(default_factory=DOCXTypography)
    toc_sections: List[str]     = field(default_factory=list)
    phase_rows:   List[DOCXPhaseRow] = field(default_factory=list)
    raw_sections: Dict[str, str] = field(default_factory=dict)  # section_name → content text


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------

class DOCXTemplateParser:
    """
    Parses the WK T360 Proposal DOCX and extracts reusable design assets.

    If python-docx is not installed, returns default hard-coded assets
    (already hand-extracted from the file).
    """

    TEMPLATE_PATH = (
        Path(__file__).resolve().parent.parent
        / "private"
        / "WK T360 Proposal - September 11, 2024.docx"
    )

    # Hard-coded fallback — extracted from the real DOCX on 2026-07-13
    DEFAULT_TOC = [
        "Executive Summary",
        "Company Profile",
        "Global Leadership",
        "Project Summary",
        "Our Proposed Solution",
        "Your Investment",
        "Appendix A",
        "Appendix B",
        "Appendix C",
        "Appendix D",
    ]

    DEFAULT_PHASES = [
        DOCXPhaseRow("Phase 1", "Week 1",  "Requirement Gathering & Workflow Analysis"),
        DOCXPhaseRow("Phase 2", "Week 2",  "UI/UX Design & System Architecture"),
        DOCXPhaseRow("Phase 3", "Week 3–6","Core Application Development"),
        DOCXPhaseRow("Phase 4", "Week 7",  "Integration, Testing & QA"),
        DOCXPhaseRow("Phase 5", "Week 8",  "Deployment & User Acceptance Testing"),
        DOCXPhaseRow("Phase 6", "O&M",     "Operations & Maintenance Support"),
    ]

    def parse(self) -> DOCXTemplateAssets:
        """
        Parse the DOCX template and return extracted assets.
        Falls back to hard-coded defaults if parsing fails.
        """
        assets = DOCXTemplateAssets(
            toc_sections=list(self.DEFAULT_TOC),
            phase_rows=list(self.DEFAULT_PHASES),
        )

        if not self.TEMPLATE_PATH.exists():
            logger.warning(
                f"[DOCXParser] Template not found at {self.TEMPLATE_PATH}. "
                "Using hard-coded fallback assets."
            )
            return assets

        try:
            from docx import Document
            doc = Document(str(self.TEMPLATE_PATH))
            self._extract_toc(doc, assets)
            self._extract_phase_table(doc, assets)
            self._extract_theme_colors(doc, assets)
            self._extract_sections_content(doc, assets)
            logger.info("[DOCXParser] Template parsed successfully.")
        except ImportError:
            logger.warning(
                "[DOCXParser] python-docx not installed. Using hard-coded fallback assets. "
                "Install with: pip install python-docx>=1.1.0"
            )
        except Exception as exc:
            logger.warning(f"[DOCXParser] Parsing failed: {exc}. Using hard-coded fallback assets.")

        return assets

    # ------------------------------------------------------------------
    # Extraction helpers
    # ------------------------------------------------------------------

    def _extract_toc(self, doc: Any, assets: DOCXTemplateAssets) -> None:
        """Extract TOC from first table."""
        try:
            if doc.tables:
                toc_table = doc.tables[0]
                toc_items = []
                for row in toc_table.rows:
                    for cell in row.cells:
                        txt = cell.text.strip()
                        if txt and txt not in toc_items and not txt.startswith("Appendix"):
                            toc_items.append(txt)
                        elif txt.startswith("Appendix"):
                            toc_items.append(txt)
                if toc_items:
                    assets.toc_sections = toc_items
        except Exception as exc:
            logger.debug(f"[DOCXParser] TOC extraction failed: {exc}")

    def _extract_phase_table(self, doc: Any, assets: DOCXTemplateAssets) -> None:
        """Extract implementation phase table (table index 2 in the DOCX)."""
        try:
            if len(doc.tables) >= 3:
                phase_tbl = doc.tables[2]
                phases = []
                for i, row in enumerate(phase_tbl.rows):
                    cells = [c.text.strip() for c in row.cells]
                    if i == 0 or not cells[0]:  # skip header
                        continue
                    if len(cells) >= 3:
                        phases.append(DOCXPhaseRow(cells[0], cells[1], cells[2]))
                if phases:
                    assets.phase_rows = phases
        except Exception as exc:
            logger.debug(f"[DOCXParser] Phase table extraction failed: {exc}")

    def _extract_theme_colors(self, doc: Any, assets: DOCXTemplateAssets) -> None:
        """Extract theme colors from document XML."""
        try:
            import xml.etree.ElementTree as ET
            try:
                theme_part = doc.part.part_related_by(
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme"
                )
                root = ET.fromstring(theme_part.blob)
                ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
                # Extract accent colors in order
                accent_tags = ["a:accent1", "a:accent2", "a:accent3", "a:accent4"]
                extracted = []
                for tag in accent_tags:
                    for el in root.findall(f".//{tag}", ns):
                        for child in el:
                            val = child.get("val") or child.get("lastClr")
                            if val and len(val) == 6:
                                extracted.append(f"#{val.upper()}")
                if len(extracted) >= 2:
                    assets.colors.accent_blue = extracted[0]
                    assets.colors.secondary = extracted[1]
            except Exception:
                pass  # use defaults
        except Exception as exc:
            logger.debug(f"[DOCXParser] Theme extraction failed: {exc}")

    def _extract_sections_content(self, doc: Any, assets: DOCXTemplateAssets) -> None:
        """Extract raw text content per section."""
        try:
            current_section = "Cover"
            section_texts: Dict[str, List[str]] = {"Cover": []}

            for para in doc.paragraphs:
                txt = para.text.strip()
                if not txt:
                    continue
                # Detect section boundary by font size >= 22pt
                is_heading = False
                for run in para.runs:
                    if run.font.size and run.font.size.pt >= 20:
                        is_heading = True
                        break

                if is_heading and txt in assets.toc_sections:
                    current_section = txt
                    if current_section not in section_texts:
                        section_texts[current_section] = []
                else:
                    section_texts.setdefault(current_section, []).append(txt)

            # Join and truncate each section
            for section, lines in section_texts.items():
                assets.raw_sections[section] = "\n".join(lines)[:3000]

        except Exception as exc:
            logger.debug(f"[DOCXParser] Section content extraction failed: {exc}")


# ---------------------------------------------------------------------------
# Module-level singleton (cached after first call)
# ---------------------------------------------------------------------------

_cached_assets: Optional[DOCXTemplateAssets] = None


def get_template_assets() -> DOCXTemplateAssets:
    """Return cached DOCX template assets (parsed once on first call)."""
    global _cached_assets
    if _cached_assets is None:
        _cached_assets = DOCXTemplateParser().parse()
    return _cached_assets


if __name__ == "__main__":
    import json
    import dataclasses
    assets = get_template_assets()
    print("=== Colors ===")
    print(json.dumps(dataclasses.asdict(assets.colors), indent=2))
    print("\n=== TOC ===")
    for s in assets.toc_sections:
        print(f"  • {s}")
    print("\n=== Phase Table ===")
    for p in assets.phase_rows:
        print(f"  {p.phase} | {p.duration} | {p.focus}")
    print("\n=== Header/Footer ===")
    print(f"  Header: {assets.header.left_label}/{assets.header.left_value} | {assets.header.right_label}/{assets.header.right_value}")
    print(f"  Footer: {assets.footer.left_text} | {assets.footer.right_text}")
